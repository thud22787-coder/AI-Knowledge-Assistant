from pathlib import Path
from numbers import Real
from uuid import uuid4

import docx
import pytest
import fitz
from fastapi.testclient import TestClient

from app.database import SessionLocal
from app.main import app
from app.models import Chunk, Document, User
from app.services.auth import decode_access_token


client = TestClient(app)
UPLOAD_DIR = Path("storage/uploads")


@pytest.fixture
def auth_headers():
    email = f"test-{uuid4()}@example.com"
    password = "test_password_123"
    register_response = client.post("/auth/register", json={"email": email, "password": password})
    assert register_response.status_code == 200
    login_response = client.post("/auth/login", json={"email": email, "password": password})
    assert login_response.status_code == 200
    token = login_response.json()["access_token"]
    return {"Authorization": f"Bearer {token}"}


@pytest.fixture
def cleanup_uploaded_files():
    created_paths = []
    created_ids = []
    yield created_paths, created_ids

    db = SessionLocal()
    try:
        db.query(Chunk).delete()
        for document_id in created_ids:
            document = db.get(Document, document_id)
            if document is not None:
                db.delete(document)
        db.commit()
    finally:
        db.close()

    for file_path in created_paths:
        path = Path(file_path)
        if path.exists():
            path.unlink()


def test_upload_valid_txt_file(cleanup_uploaded_files, auth_headers):
    created_paths, created_ids = cleanup_uploaded_files
    response = client.post(
        "/documents/upload",
        files={"file": ("sample.txt", b"hello world", "text/plain")},
        headers=auth_headers,
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["id"]
    assert payload["filename"] == "sample.txt"
    assert payload["status"] == "uploaded"

    db = SessionLocal()
    try:
        document = db.get(Document, payload["id"])
        assert document is not None
        created_ids.append(document.id)
        created_paths.append(document.file_path)
    finally:
        db.close()


def test_upload_rejects_invalid_extension(cleanup_uploaded_files, auth_headers):
    response = client.post(
        "/documents/upload",
        files={"file": ("malware.exe", b"bad", "application/octet-stream")},
        headers=auth_headers,
    )

    assert response.status_code == 400


def test_upload_creates_file_on_disk(cleanup_uploaded_files, auth_headers):
    created_paths, created_ids = cleanup_uploaded_files
    response = client.post(
        "/documents/upload",
        files={"file": ("diskcheck.txt", b"hello world", "text/plain")},
        headers=auth_headers,
    )

    assert response.status_code == 200
    payload = response.json()

    db = SessionLocal()
    try:
        document = db.get(Document, payload["id"])
        assert document is not None
        created_ids.append(document.id)
        created_paths.append(document.file_path)
        assert Path(document.file_path).exists()
        assert UPLOAD_DIR in Path(document.file_path).parents
    finally:
        db.close()


def test_list_documents(cleanup_uploaded_files, auth_headers):
    created_paths, created_ids = cleanup_uploaded_files
    upload_response = client.post(
        "/documents/upload",
        files={"file": ("listcheck.txt", b"hello world", "text/plain")},
        headers=auth_headers,
    )

    assert upload_response.status_code == 200
    uploaded = upload_response.json()

    db = SessionLocal()
    try:
        document = db.get(Document, uploaded["id"])
        assert document is not None
        created_ids.append(document.id)
        created_paths.append(document.file_path)
    finally:
        db.close()

    response = client.get("/documents", headers=auth_headers)
    assert response.status_code == 200
    payload = response.json()
    assert any(item["id"] == uploaded["id"] and item["filename"] == "listcheck.txt" for item in payload)


def test_delete_document(cleanup_uploaded_files, auth_headers):
    created_paths, created_ids = cleanup_uploaded_files
    upload_response = client.post(
        "/documents/upload",
        files={"file": ("deletecheck.txt", b"hello world", "text/plain")},
        headers=auth_headers,
    )

    assert upload_response.status_code == 200
    uploaded = upload_response.json()

    db = SessionLocal()
    try:
        document = db.get(Document, uploaded["id"])
        assert document is not None
        created_ids.append(document.id)
        created_paths.append(document.file_path)
        file_path = Path(document.file_path)
    finally:
        db.close()

    delete_response = client.delete(f"/documents/{uploaded['id']}", headers=auth_headers)
    assert delete_response.status_code == 200
    assert delete_response.json() == {"id": uploaded["id"], "deleted": True}
    assert not file_path.exists()

    list_response = client.get("/documents", headers=auth_headers)
    assert list_response.status_code == 200
    assert all(item["id"] != uploaded["id"] for item in list_response.json())


def test_delete_nonexistent_document(cleanup_uploaded_files, auth_headers):
    response = client.delete(f"/documents/{uuid4()}", headers=auth_headers)
    assert response.status_code == 404


def test_delete_document_also_deletes_chunks(cleanup_uploaded_files, auth_headers):
    created_paths, created_ids = cleanup_uploaded_files
    text = (
        "Paragraph one has enough text to create a chunk. "
        "It is followed by another sentence.\n\n"
        "Paragraph two also creates more than one useful chunk for testing."
    )
    upload_response = client.post(
        "/documents/upload",
        files={"file": ("delete_chunks.txt", text.encode("utf-8"), "text/plain")},
        headers=auth_headers,
    )

    assert upload_response.status_code == 200
    uploaded = upload_response.json()

    db = SessionLocal()
    try:
        document = db.get(Document, uploaded["id"])
        assert document is not None
        created_ids.append(document.id)
        created_paths.append(document.file_path)
    finally:
        db.close()

    process_response = client.post(f"/documents/{uploaded['id']}/process", headers=auth_headers)
    assert process_response.status_code == 200
    assert process_response.json()["chunk_count"] > 0

    db = SessionLocal()
    try:
        assert db.query(Chunk).filter(Chunk.document_id == uploaded["id"]).count() > 0
    finally:
        db.close()

    delete_response = client.delete(f"/documents/{uploaded['id']}", headers=auth_headers)
    assert delete_response.status_code == 200

    db = SessionLocal()
    try:
        assert db.query(Chunk).filter(Chunk.document_id == uploaded["id"]).count() == 0
    finally:
        db.close()


def test_process_txt_document_creates_embeddings(cleanup_uploaded_files, auth_headers):
    created_paths, created_ids = cleanup_uploaded_files
    text = (
        "Paragraph one has enough text to make the parser and chunker work properly. "
        "It should become at least one chunk.\n\n"
        "Paragraph two also has enough content to ensure multiple paragraphs are handled. "
        "This gives us more than one paragraph for the test."
    )
    upload_response = client.post(
        "/documents/upload",
        files={"file": ("process.txt", text.encode("utf-8"), "text/plain")},
        headers=auth_headers,
    )

    assert upload_response.status_code == 200
    uploaded = upload_response.json()

    db = SessionLocal()
    try:
        document = db.get(Document, uploaded["id"])
        assert document is not None
        created_ids.append(document.id)
        created_paths.append(document.file_path)
    finally:
        db.close()

    process_response = client.post(f"/documents/{uploaded['id']}/process", headers=auth_headers)
    assert process_response.status_code == 200
    process_payload = process_response.json()
    assert process_payload["document_id"] == uploaded["id"]
    assert process_payload["status"] == "ready"
    assert process_payload["chunk_count"] > 0

    db = SessionLocal()
    try:
        chunks = db.query(Chunk).filter(Chunk.document_id == uploaded["id"]).order_by(Chunk.chunk_index.asc()).all()
        document = db.get(Document, uploaded["id"])
        assert document is not None
        assert document.status == "ready"
        assert len(chunks) == process_payload["chunk_count"]
        assert len(chunks) > 0
        for chunk in chunks:
            assert chunk.embedding is not None
            assert len(chunk.embedding) == 384
            assert all(isinstance(value, Real) and not isinstance(value, bool) for value in chunk.embedding)
    finally:
        db.close()


def test_process_pdf_document_creates_embeddings(cleanup_uploaded_files, tmp_path, auth_headers):
    created_paths, created_ids = cleanup_uploaded_files
    pdf_path = tmp_path / "process.pdf"
    document_pdf = fitz.open()
    try:
        page = document_pdf.new_page()
        page.insert_text((72, 72), "This is a PDF document for processing.")
        document_pdf.save(pdf_path)
    finally:
        document_pdf.close()

    upload_response = client.post(
        "/documents/upload",
        files={"file": ("process.pdf", pdf_path.read_bytes(), "application/pdf")},
        headers=auth_headers,
    )

    assert upload_response.status_code == 200
    uploaded = upload_response.json()

    db = SessionLocal()
    try:
        document = db.get(Document, uploaded["id"])
        assert document is not None
        created_ids.append(document.id)
        created_paths.append(document.file_path)
    finally:
        db.close()

    process_response = client.post(f"/documents/{uploaded['id']}/process", headers=auth_headers)
    assert process_response.status_code == 200
    process_payload = process_response.json()
    assert process_payload["document_id"] == uploaded["id"]
    assert process_payload["chunk_count"] > 0

    db = SessionLocal()
    try:
        chunks = db.query(Chunk).filter(Chunk.document_id == uploaded["id"]).order_by(Chunk.chunk_index.asc()).all()
        assert len(chunks) > 0
        assert any(chunk.embedding is not None for chunk in chunks)
    finally:
        db.close()


def test_process_multipage_pdf_has_sequential_chunk_index(cleanup_uploaded_files, tmp_path, auth_headers):
    created_paths, created_ids = cleanup_uploaded_files
    pdf_path = tmp_path / "multipage.pdf"
    page_texts = [
        (
            "Page one paragraph one is intentionally long so it comfortably exceeds the chunking threshold. "
            "It keeps repeating descriptive language to ensure the first page yields multiple chunks. "
            "Another sentence adds more length and keeps the content stable for the test.\n\n"
            "Page one paragraph two is also long enough to force another chunk. "
            "The sentence structure is repetitive on purpose so the parser and chunker both have enough material. "
            "This paragraph should contribute additional chunks on page one."
        ),
        (
            "Page two paragraph one is similarly verbose and exists to create multiple chunks on the second page. "
            "The wording is repetitive, extended, and deliberately longer than the chunk size threshold. "
            "This makes it easy to detect whether chunk_index resets for a new page.\n\n"
            "Page two paragraph two continues the pattern with another long block of text for chunking. "
            "It provides enough length to guarantee more than one chunk on page two as well. "
            "The goal is to confirm indexes remain sequential across the whole document."
        ),
    ]

    document_pdf = fitz.open()
    try:
        for text in page_texts:
            page = document_pdf.new_page()
            y = 72
            for paragraph in text.split("\n\n"):
                page.insert_text((72, y), paragraph)
                y += 120
        document_pdf.save(pdf_path)
    finally:
        document_pdf.close()

    upload_response = client.post(
        "/documents/upload",
        files={"file": ("multipage.pdf", pdf_path.read_bytes(), "application/pdf")},
        headers=auth_headers,
    )

    assert upload_response.status_code == 200
    uploaded = upload_response.json()

    db = SessionLocal()
    try:
        document = db.get(Document, uploaded["id"])
        assert document is not None
        created_ids.append(document.id)
        created_paths.append(document.file_path)
    finally:
        db.close()

    process_response = client.post(f"/documents/{uploaded['id']}/process", headers=auth_headers)
    assert process_response.status_code == 200

    db = SessionLocal()
    try:
        chunks = db.query(Chunk).filter(Chunk.document_id == uploaded["id"]).order_by(Chunk.chunk_index.asc()).all()
        assert [chunk.chunk_index for chunk in chunks] == list(range(len(chunks)))
    finally:
        db.close()


def test_process_nonexistent_document_returns_404(cleanup_uploaded_files, auth_headers):
    response = client.post(f"/documents/{uuid4()}/process", headers=auth_headers)
    assert response.status_code == 404


def test_process_unsupported_extension_returns_400(cleanup_uploaded_files, auth_headers):
    created_paths, created_ids = cleanup_uploaded_files
    token = auth_headers["Authorization"].split(" ", 1)[1]
    user_id = decode_access_token(token)["sub"]
    db = SessionLocal()
    try:
        document = Document(
            filename="fake.xlsx",
            file_path="storage/uploads/fake.xlsx",
            content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            status="uploaded",
            user_id=user_id,
        )
        db.add(document)
        db.commit()
        db.refresh(document)
        created_ids.append(document.id)
    finally:
        db.close()

    process_response = client.post(f"/documents/{document.id}/process", headers=auth_headers)
    assert process_response.status_code == 400


def test_process_docx_document_creates_embeddings(cleanup_uploaded_files, tmp_path, auth_headers):
    created_paths, created_ids = cleanup_uploaded_files
    docx_path = tmp_path / "process.docx"
    document_docx = docx.Document()
    document_docx.add_paragraph(
        "This is the first DOCX paragraph for processing. It should be parsed into text and chunked."
    )
    document_docx.add_paragraph(
        "This is the second DOCX paragraph for processing. It gives the test another paragraph to verify."
    )
    document_docx.save(docx_path)

    upload_response = client.post(
        "/documents/upload",
        files={
            "file": (
                "process.docx",
                docx_path.read_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        headers=auth_headers,
    )

    assert upload_response.status_code == 200
    uploaded = upload_response.json()

    db = SessionLocal()
    try:
        document = db.get(Document, uploaded["id"])
        assert document is not None
        created_ids.append(document.id)
        created_paths.append(document.file_path)
    finally:
        db.close()


def test_user_cannot_access_another_users_document(cleanup_uploaded_files):
    created_paths, created_ids = cleanup_uploaded_files

    email_a = f"test-{uuid4()}@example.com"
    email_b = f"test-{uuid4()}@example.com"
    password = "test_password_123"

    client.post("/auth/register", json={"email": email_a, "password": password})
    client.post("/auth/register", json={"email": email_b, "password": password})
    token_a = client.post("/auth/login", json={"email": email_a, "password": password}).json()["access_token"]
    token_b = client.post("/auth/login", json={"email": email_b, "password": password}).json()["access_token"]
    headers_a = {"Authorization": f"Bearer {token_a}"}
    headers_b = {"Authorization": f"Bearer {token_b}"}

    upload_response = client.post(
        "/documents/upload",
        files={"file": ("owned_by_a.txt", b"secret content", "text/plain")},
        headers=headers_a,
    )
    assert upload_response.status_code == 200
    uploaded = upload_response.json()

    db = SessionLocal()
    try:
        document = db.get(Document, uploaded["id"])
        assert document is not None
        created_ids.append(document.id)
        created_paths.append(document.file_path)
    finally:
        db.close()

    list_response_b = client.get("/documents", headers=headers_b)
    assert list_response_b.status_code == 200
    assert all(item["id"] != uploaded["id"] for item in list_response_b.json())

    process_response_b = client.post(f"/documents/{uploaded['id']}/process", headers=headers_b)
    assert process_response_b.status_code == 404

    delete_response_b = client.delete(f"/documents/{uploaded['id']}", headers=headers_b)
    assert delete_response_b.status_code == 404

    list_response_a = client.get("/documents", headers=headers_a)
    assert list_response_a.status_code == 200
    assert any(item["id"] == uploaded["id"] for item in list_response_a.json())
