import json
import os
from pathlib import Path
from tempfile import TemporaryDirectory
from unittest.mock import patch
from uuid import uuid4

import docx
import fitz
from fastapi.testclient import TestClient

from app.database import SessionLocal
from app.main import app
from app.services.auth import decode_access_token
from app.services.embedder import embed_text
from app.services.retriever import search_similar_chunks
from eval.golden_dataset import GOLDEN_DATASET


client = TestClient(app)


def register_and_login() -> tuple[dict[str, str], str]:
    email = f"eval-{uuid4()}@example.com"
    password = "test_password_123"
    register_response = client.post("/auth/register", json={"email": email, "password": password})
    register_response.raise_for_status()
    login_response = client.post("/auth/login", json={"email": email, "password": password})
    login_response.raise_for_status()
    token = login_response.json()["access_token"]
    user_id = decode_access_token(token)["sub"]
    return {"Authorization": f"Bearer {token}"}, user_id


def create_sample_files(base_dir: Path) -> list[tuple[str, Path, str]]:
    files: list[tuple[str, Path, str]] = []

    sample_pdf = base_dir / "sample.pdf"
    pdf_doc = fitz.open()
    try:
        for text in [
            "Page one: the quick brown fox jumps over the lazy dog.",
            "Page two: pack my box with five dozen liquor jugs.",
        ]:
            page = pdf_doc.new_page()
            page.insert_text((72, 72), text)
        pdf_doc.save(sample_pdf)
    finally:
        pdf_doc.close()
    files.append(("sample.pdf", sample_pdf, "application/pdf"))

    sample_docx = base_dir / "sample.docx"
    doc = docx.Document()
    for text in [
        "DOCX paragraph one: alpha beta gamma.",
        "DOCX paragraph two: delta epsilon zeta.",
        "DOCX paragraph three: eta theta iota.",
    ]:
        doc.add_paragraph(text)
    doc.save(sample_docx)
    files.append(
        (
            "sample.docx",
            sample_docx,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    )

    include_process_txt = os.getenv("RAW_CHAT_RUN_INCLUDE_PROCESS_TXT", "1") == "1"
    if include_process_txt:
        process_txt = base_dir / "process.txt"
        process_txt.write_text(
            "Paragraph one has enough text to make the parser and chunker work properly. "
            "It should become at least one chunk.\n\n"
            "Paragraph two also has enough content to ensure multiple paragraphs are handled. "
            "This gives us more than one paragraph for the test.",
            encoding="utf-8",
        )
        files.append(("process.txt", process_txt, "text/plain"))

    process_pdf = base_dir / "process.pdf"
    pdf_doc = fitz.open()
    try:
        page = pdf_doc.new_page()
        page.insert_text((72, 72), "This is a PDF document for processing.")
        pdf_doc.save(process_pdf)
    finally:
        pdf_doc.close()
    files.append(("process.pdf", process_pdf, "application/pdf"))

    process_docx = base_dir / "process.docx"
    doc = docx.Document()
    doc.add_paragraph("This is the first DOCX paragraph for processing. It should be parsed into text and chunked.")
    doc.add_paragraph("This is the second DOCX paragraph for processing. It gives the test another paragraph to verify.")
    doc.save(process_docx)
    files.append(
        (
            "process.docx",
            process_docx,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    )

    return files


def upload_and_process(headers: dict[str, str], files: list[tuple[str, Path, str]]) -> None:
    for filename, path, content_type in files:
        upload_response = client.post(
            "/documents/upload",
            files={"file": (filename, path.read_bytes(), content_type)},
            headers=headers,
        )
        upload_response.raise_for_status()
        document_id = upload_response.json()["id"]

        process_response = client.post(f"/documents/{document_id}/process", headers=headers)
        process_response.raise_for_status()


def fake_generate_answer(question: str, context_chunks: list, history=None) -> str:
    if not context_chunks:
        return "Không tìm thấy thông tin liên quan trong tài liệu đã upload."
    return context_chunks[0].text


def run_eval() -> list[dict]:
    headers, user_id = register_and_login()

    with TemporaryDirectory() as temp_dir:
        sample_files = create_sample_files(Path(temp_dir))
        upload_and_process(headers, sample_files)

    results: list[dict] = []

    with patch("app.routers.chat.generate_answer", side_effect=fake_generate_answer):
        for item in GOLDEN_DATASET:
            question = item["question"]

            db = SessionLocal()
            try:
                query_embedding = embed_text(question)
                retrieved_chunks = search_similar_chunks(
                    query_embedding=query_embedding,
                    user_id=user_id,
                    db=db,
                    top_k=20,
                )
                retrieved_contexts = [chunk.text for chunk in retrieved_chunks]
            finally:
                db.close()

            chat_response = client.post("/chat", json={"question": question}, headers=headers)
            chat_response.raise_for_status()
            chat_payload = chat_response.json()

            results.append(
                {
                    "question": question,
                    "ground_truth": item["ground_truth"],
                    "source_doc": item["source_doc"],
                    "answer": chat_payload["answer"],
                    "retrieved_contexts": retrieved_contexts,
                    "reranked_sources": chat_payload["sources"],
                }
            )

    return results


def main() -> None:
    output_path = Path("eval/raw_run_output.json")
    results = run_eval()
    output_path.write_text(json.dumps(results, ensure_ascii=False, indent=2), encoding="utf-8")
    print(output_path.read_text(encoding="utf-8"))


if __name__ == "__main__":
    main()
