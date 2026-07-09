from __future__ import annotations

import asyncio
import json
import sys
from pathlib import Path
from tempfile import TemporaryDirectory
from uuid import uuid4

import docx
import fitz
from fastapi.testclient import TestClient
from openai import AsyncOpenAI
from ragas.embeddings import OpenAIEmbeddings
from ragas.llms import llm_factory
from ragas.metrics.collections import AnswerRelevancy, ContextPrecision, ContextRecall, Faithfulness

ROOT_DIR = Path(__file__).resolve().parents[1]
if str(ROOT_DIR) not in sys.path:
    sys.path.insert(0, str(ROOT_DIR))

from app.main import app
from eval.golden_dataset import GOLDEN_DATASET


client = TestClient(app)


def register_and_login() -> dict[str, str]:
    email = f"ragas-eval-{uuid4()}@example.com"
    password = "test_password_123"

    register_response = client.post("/auth/register", json={"email": email, "password": password})
    register_response.raise_for_status()

    login_response = client.post("/auth/login", json={"email": email, "password": password})
    login_response.raise_for_status()

    token = login_response.json()["access_token"]
    return {"Authorization": f"Bearer {token}"}


def create_test_files(base_dir: Path) -> list[Path]:
    files: list[Path] = []

    process_txt = base_dir / "process.txt"
    process_txt.write_text(
        "Paragraph one has enough text to make the parser and chunker work properly. "
        "It should become at least one chunk.\n\n"
        "Paragraph two also has enough content to ensure multiple paragraphs are handled. "
        "This gives us more than one paragraph for the test.",
        encoding="utf-8",
    )
    files.append(process_txt)

    process_pdf = base_dir / "process.pdf"
    process_pdf_doc = fitz.open()
    try:
        page = process_pdf_doc.new_page()
        page.insert_text((72, 72), "This is a PDF document for processing.")
        process_pdf_doc.save(process_pdf)
    finally:
        process_pdf_doc.close()
    files.append(process_pdf)

    process_docx = base_dir / "process.docx"
    process_docx_doc = docx.Document()
    process_docx_doc.add_paragraph(
        "This is the first DOCX paragraph for processing. It should be parsed into text and chunked."
    )
    process_docx_doc.add_paragraph(
        "This is the second DOCX paragraph for processing. It gives the test another paragraph to verify."
    )
    process_docx_doc.save(process_docx)
    files.append(process_docx)

    sample_pdf = base_dir / "sample.pdf"
    sample_pdf_doc = fitz.open()
    try:
        for text in [
            "Page one: the quick brown fox jumps over the lazy dog.",
            "Page two: pack my box with five dozen liquor jugs.",
        ]:
            page = sample_pdf_doc.new_page()
            page.insert_text((72, 72), text)
        sample_pdf_doc.save(sample_pdf)
    finally:
        sample_pdf_doc.close()
    files.append(sample_pdf)

    sample_docx = base_dir / "sample.docx"
    sample_docx_doc = docx.Document()
    for text in [
        "DOCX paragraph one: alpha beta gamma.",
        "DOCX paragraph two: delta epsilon zeta.",
        "DOCX paragraph three: eta theta iota.",
    ]:
        sample_docx_doc.add_paragraph(text)
    sample_docx_doc.save(sample_docx)
    files.append(sample_docx)

    return files


def upload_and_process(file_path: Path, headers: dict[str, str]) -> None:
    content_type_by_suffix = {
        ".txt": "text/plain",
        ".pdf": "application/pdf",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    }
    with file_path.open("rb") as file_obj:
        upload_response = client.post(
            "/documents/upload",
            files={"file": (file_path.name, file_obj, content_type_by_suffix[file_path.suffix.lower()])},
            headers=headers,
        )
    upload_response.raise_for_status()

    document_id = upload_response.json()["id"]
    process_response = client.post(f"/documents/{document_id}/process", headers=headers)
    process_response.raise_for_status()


def run() -> list[dict[str, object]]:
    headers = register_and_login()

    with TemporaryDirectory() as temp_dir:
        temp_path = Path(temp_dir)
        for file_path in create_test_files(temp_path):
            upload_and_process(file_path, headers)

    results: list[dict[str, object]] = []
    for item in GOLDEN_DATASET:
        chat_response = client.post("/chat", json={"question": item["question"]}, headers=headers)
        chat_response.raise_for_status()
        payload = chat_response.json()
        results.append(
            {
                "question": item["question"],
                "answer": payload["answer"],
                "contexts": [source["text"] for source in payload["sources"]],
                "ground_truth": item["ground_truth"],
            }
        )

    return results


async def run_ragas_metrics(results: list[dict[str, object]]) -> dict[str, float]:
    client = AsyncOpenAI()
    llm = llm_factory("gpt-4o-mini", client=client)
    embeddings = OpenAIEmbeddings(client=client)

    faithfulness_metric = Faithfulness(llm=llm)
    answer_relevancy_metric = AnswerRelevancy(llm=llm, embeddings=embeddings)
    context_precision_metric = ContextPrecision(llm=llm)
    context_recall_metric = ContextRecall(llm=llm)

    scores: dict[str, list[float]] = {
        "faithfulness": [],
        "answer_relevancy": [],
        "context_precision": [],
        "context_recall": [],
    }

    for item in results:
        faithfulness_result = await faithfulness_metric.ascore(
            user_input=item["question"],
            response=item["answer"],
            retrieved_contexts=item["contexts"],
        )
        scores["faithfulness"].append(faithfulness_result.value)

        answer_relevancy_result = await answer_relevancy_metric.ascore(
            user_input=item["question"],
            response=item["answer"],
        )
        scores["answer_relevancy"].append(answer_relevancy_result.value)

        context_precision_result = await context_precision_metric.ascore(
            user_input=item["question"],
            reference=item["ground_truth"],
            retrieved_contexts=item["contexts"],
        )
        scores["context_precision"].append(context_precision_result.value)

        context_recall_result = await context_recall_metric.ascore(
            user_input=item["question"],
            retrieved_contexts=item["contexts"],
            reference=item["ground_truth"],
        )
        scores["context_recall"].append(context_recall_result.value)

    return {name: sum(values) / len(values) for name, values in scores.items() if values}


if __name__ == "__main__":
    results = run()
    output_path = Path("eval/results.json")
    output_path.write_text(json.dumps(results, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"Wrote {len(results)} results to eval/results.json")

    average_scores = asyncio.run(run_ragas_metrics(results))
    scores_path = Path("eval/ragas_scores.json")
    scores_path.write_text(json.dumps(average_scores, ensure_ascii=False, indent=2), encoding="utf-8")
    print("RAGAS average scores:", average_scores)
